import streamlit as st
import pandas as pd
import base64
from docx import Document
import util  # make sure util.py has create_docx() and create_combined_docx()

# ----------------------------
# Page configuration
# ----------------------------
st.set_page_config(page_title="Project Summary Generator", layout="wide")
st.title("Project Summary Generator")

# ----------------------------
# File Upload Section
# ----------------------------
uploaded_file = st.file_uploader(
    "Upload your data file (.xlsx, .csv, .txt)",
    type=["xlsx", "csv", "txt"]
)

# ----------------------------
# File Reading Logic
# ----------------------------
if uploaded_file is not None:
    file_ready = False
    try:
        if uploaded_file.name.endswith('.xlsx'):
            df = pd.read_excel(uploaded_file, engine='openpyxl')
            file_ready = True
        elif uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
            file_ready = True
        elif uploaded_file.name.endswith('.txt'):
            df = pd.read_csv(uploaded_file, delimiter="\t")
            file_ready = True
        else:
            st.error("❌ Unsupported file format.")
            st.stop()
    except Exception as e:
        st.error(f"⚠️ Could not read the uploaded file: {e}")
        st.stop()

    # ----------------------------
    # Validation and Warnings
    # ----------------------------
    if file_ready:
        required_cols = [
            'p_number', 'short_description', 'description',
            'affected_customers', 'state', 'completion_code'
        ]
        missing = [c for c in required_cols if c not in df.columns]
        if missing:
            st.warning(
                f"⚠️ Uploaded file is missing expected columns: {missing}. "
                "The app will attempt to continue but results may be incomplete."
            )

        # Accessible high-contrast CSS
        st.markdown("""
            <style>
            /* Page container and block background: keep light and neutral for readability */
            .main, .block-container {
                background: #ffffff !important;
                color: #0b1a2b !important; /* very dark blue/near-black for text */
            }

            /* Make the app container slightly elevated */
            .block-container {
                padding-top: 1.5rem;
                padding-bottom: 1.5rem;
                border-radius: 10px;
                box-shadow: 0 2px 8px rgba(11,26,43,0.06);
            }

            /* Primary buttons: dark blue background with white text */
            .stButton>button, .stDownloadButton>button {
                background-color: #0b69ff !important; /* accessible blue */
                color: #ffffff !important;
                border-radius: 6px !important;
                font-weight: 600 !important;
                padding: 0.45em 1.1em !important;
                border: 1px solid rgba(11,105,255,0.12) !important;
            }

            /* Secondary download buttons (if any) use green with white text */
            .stDownloadButton>button {
                background-color: #047857 !important; /* emerald green */
                border: 1px solid rgba(4,120,87,0.12) !important;
            }

            /* Text areas and inputs: white background, dark text, clear border */
            .stTextArea textarea, .stTextInput>div>input, .stSelectbox>div>div>div {
                background: #ffffff !important;
                color: #0b1a2b !important;
                border: 1px solid #cbd5e1 !important;
                border-radius: 6px !important;
            }

            /* File uploader look */
            .stFileUploader>div, .stFileUploader {
                background: #f8fafc !important;
                color: #0b1a2b !important;
                border: 1px dashed #cbd5e1 !important;
                border-radius: 6px !important;
            }

            /* Make headings and project identifiers prominent */
            h1, h2, h3 {
                color: #07122a !important;
            }

            /* Links and download links should be clearly visible */
            a { color: #0b69ff !important; font-weight:600; }

            /* Captions and small notes: slightly muted but still readable */
            .stCaption, .caption {
                color: #334155 !important;
            }

            /* Tables and markdown text */
            .stMarkdown, .stTable, .stDataFrame, .stWrite {
                color: #0b1a2b !important;
            }

            /* Ensure contrast for text areas used as preview */
            .stTextArea textarea { height: 300px !important; }
            </style>
        """, unsafe_allow_html=True)

        # Sidebar for instructions and branding
        st.sidebar.image("https://img.icons8.com/color/96/000000/document--v2.png", width=64)
        st.sidebar.title("Project Summary Generator")
        st.sidebar.markdown("""
        **Instructions:**
        1. Upload your data file (.xlsx, .csv, .txt)
        2. Click **Generate** to view and download project summaries
        3. Use **Preview** to read and understand each summary before downloading
        """)
        st.sidebar.markdown("---")
        st.sidebar.markdown("Made with ❤️ by Your Team")

        # ----------------------------
        # Generate Button Logic (uses session_state)
        # ----------------------------
        if st.button("Generate Summaries"):
            st.session_state['generated'] = True
            st.session_state['preview_index'] = None

        # ----------------------------
        # If summaries are generated
        # ----------------------------
        if st.session_state.get('generated', False):
            # Download all summaries
            combined_buf = util.create_combined_docx(df)
            st.download_button(
                label="📥 Download All Summaries (.docx)",
                data=combined_buf.getvalue(),
                file_name="All_Project_Summaries.docx",
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            st.markdown("---")
            st.write("## 📋 Projects Table")

            preview_index = st.session_state.get('preview_index', None)

            # ----------------------------
            # Display each project in a table-like format
            # ----------------------------
            for idx, row in df.iterrows():
                title = row.get('short_description') or 'No title provided'
                pnum = row.get('p_number') or f'row-{idx}'
                doc_buf = util.create_docx(row)
                b64 = base64.b64encode(doc_buf.getvalue()).decode()
                download_link = (
                    f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;'
                    f'base64,{b64}" download="Project_{pnum}.docx">Download Summary</a>'
                )

                col1, col2, col3, col4 = st.columns([2, 4, 3, 2])
                with col1:
                    st.markdown(f"<span style='color:#111827; font-size:1.2em; font-weight:700;'>{pnum}</span>", unsafe_allow_html=True)
                with col2:
                    st.write(title)
                with col3:
                    st.markdown(download_link, unsafe_allow_html=True)
                with col4:
                    if st.button("Preview", key=f"preview_btn_{idx}"):
                        st.session_state['preview_index'] = idx

            # ----------------------------
            # Show preview for selected project
            # ----------------------------
            if st.session_state.get('preview_index', None) is not None:
                idx = st.session_state['preview_index']
                row = df.iloc[idx]
                st.markdown(f"### 🔍 Preview: Project {row.get('p_number', idx)}")

                doc_buf = util.create_docx(row)
                st.download_button(
                    label=f"Download Project {row.get('p_number', idx)} (.docx)",
                    data=doc_buf.getvalue(),
                    file_name=f"Project_{row.get('p_number', idx)}.docx",
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

                # Extract text for preview
                doc_buf.seek(0)
                doc = Document(doc_buf)
                preview_text = ""

                for para in doc.paragraphs:
                    preview_text += para.text + "\n"
                for table in doc.tables:
                    for row_cells in table.rows:
                        preview_text += " | ".join(cell.text for cell in row_cells.cells) + "\n"

                st.text_area("Generated Summary Preview", preview_text, height=350)
