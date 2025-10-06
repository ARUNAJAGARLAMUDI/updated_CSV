
import io
from docx import Document

def create_docx(row):
    """Creates a .docx file in memory for a single project."""
    document = Document()
    
    p_num = row.get('p_number', 'N/A')
    short_desc = row.get('short_description', 'No title provided')
    
    document.add_heading(f"Project Summary: {p_num}", level=1)
    document.add_heading(short_desc, level=2)

    # Using a table for better formatting
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Details'

    details = {
        "Project Number": p_num,
        "Description": row.get('description', 'Not available'),
        "Affected Customers": str(row.get('affected_customers', 'Not available')),
        "State": row.get('state', 'Not available'),
        "Completion Code": row.get('completion_code', 'Not available')
    }

    for key, value in details.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value

    # Add technical term explanations for laypersons
    document.add_heading("Technical Terms Explained", level=2)
    explanations = [
        {
            "term": "Completion Code",
            "definition": "A code that indicates the status or result of a project or task.",
            "explanation": "It helps track whether a project is finished, pending, or needs further action.",
            "example": "For example, a completion code of 'DONE' means the project is finished, while 'PENDING' means it is still in progress."
        },
        {
            "term": "Affected Customers",
            "definition": "The people or organizations impacted by the project or issue.",
            "explanation": "This term shows who will benefit from or be influenced by the project's outcome.",
            "example": "For example, if a software update fixes a bug, the affected customers are those who use that software."
        },
        {
            "term": "State",
            "definition": "The current condition or phase of the project.",
            "explanation": "It tells you if the project is new, ongoing, completed, or on hold.",
            "example": "For example, a project in the 'Active' state is currently being worked on, while 'Closed' means it is finished."
        },
        {
            "term": "Description",
            "definition": "A detailed explanation of the project or issue.",
            "explanation": "It helps everyone understand what the project is about and what it aims to achieve.",
            "example": "For example, a description might say: 'This project upgrades the company website to improve speed and security.'"
        }
    ]
    for item in explanations:
        document.add_heading(item["term"], level=3)
        document.add_paragraph(f"Definition: {item['definition']}")
        document.add_paragraph(f"Explanation: {item['explanation']}")
        document.add_paragraph(f"Example: {item['example']}")

    # Save document to an in-memory buffer
    doc_buf = io.BytesIO()
    document.save(doc_buf)
    doc_buf.seek(0)
    return doc_buf

def create_combined_docx(df):
    """Creates a single .docx file in memory containing all project summaries."""
    combined_document = Document()
    combined_document.add_heading("All Project Summaries", level=0)


    for idx, row in df.iterrows():
        p_num = row.get('p_number', 'N/A')
        short_desc = row.get('short_description', 'No title provided')

        combined_document.add_heading(f"Project Summary: {p_num}", level=1)
        combined_document.add_heading(short_desc, level=2)

        table = combined_document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Details'

        details = {
            "Project Number": p_num,
            "Description": row.get('description', 'Not available'),
            "Affected Customers": str(row.get('affected_customers', 'Not available')),
            "State": row.get('state', 'Not available'),
            "Completion Code": row.get('completion_code', 'Not available')
        }

        for key, value in details.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value

        # Add technical term explanations for laypersons
        combined_document.add_heading("Technical Terms Explained", level=2)
        explanations = [
            {
                "term": "Completion Code",
                "definition": "A code that indicates the status or result of a project or task.",
                "explanation": "It helps track whether a project is finished, pending, or needs further action.",
                "example": "For example, a completion code of 'DONE' means the project is finished, while 'PENDING' means it is still in progress."
            },
            {
                "term": "Affected Customers",
                "definition": "The people or organizations impacted by the project or issue.",
                "explanation": "This term shows who will benefit from or be influenced by the project's outcome.",
                "example": "For example, if a software update fixes a bug, the affected customers are those who use that software."
            },
            {
                "term": "State",
                "definition": "The current condition or phase of the project.",
                "explanation": "It tells you if the project is new, ongoing, completed, or on hold.",
                "example": "For example, a project in the 'Active' state is currently being worked on, while 'Closed' means it is finished."
            },
            {
                "term": "Description",
                "definition": "A detailed explanation of the project or issue.",
                "explanation": "It helps everyone understand what the project is about and what it aims to achieve.",
                "example": "For example, a description might say: 'This project upgrades the company website to improve speed and security.'"
            }
        ]
        for item in explanations:
            combined_document.add_heading(item["term"], level=3)
            combined_document.add_paragraph(f"Definition: {item['definition']}")
            combined_document.add_paragraph(f"Explanation: {item['explanation']}")
            combined_document.add_paragraph(f"Example: {item['example']}")

        # Add a page break after each project summary, except for the last one
        if idx < len(df) - 1:
            combined_document.add_page_break()

    # Save combined document to an in-memory buffer
    combined_buf = io.BytesIO()
    combined_document.save(combined_buf)
    combined_buf.seek(0)
    return combined_buf
